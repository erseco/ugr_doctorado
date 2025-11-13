import argparse
import os
import tempfile
from typing import List

from docx import Document
from gtts import gTTS
from pydub import AudioSegment


def extract_text_from_docx(path: str) -> str:
    """
    Extrai texto de parágrafos e células de tabelas de um .docx.
    Ignora imagens e objetos incorporados.
    """
    doc = Document(path)
    chunks = []

    # Parágrafos
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)

    # Tabelas (linhas -> células)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                chunks.append(" | ".join(row_text))

    return "\n".join(chunks)


def sentencewise_chunks(text: str, max_chars: int = 2000) -> List[str]:
    """
    Divide o texto em blocos <= max_chars, preservando quebras por sentença/parágrafo sempre que possível.
    """
    if not text:
        return []

    # Primeiro, quebre por duplas de quebras de linha (parágrafos maiores)
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    blocks = []
    current = ""

    def flush_current():
        nonlocal current
        if current.strip():
            blocks.append(current.strip())
        current = ""

    for p in paragraphs:
        # Se o parágrafo couber inteiro, tenta anexar ao bloco atual, senão quebra por sentenças
        if len(p) <= max_chars:
            if len(current) + len(p) + 1 <= max_chars:
                current = (current + "\n" + p).strip()
            else:
                flush_current()
                current = p
        else:
            # Parágrafo muito grande: quebre por sentenças simplificadas
            sentences = []
            buf = ""
            for token in p.replace("!", ".").replace("?", ".").split("."):
                token = token.strip()
                if not token:
                    continue
                candidate = (buf + " " + token + ".").strip()
                if len(candidate) > max_chars:
                    if buf:
                        sentences.append(buf.strip())
                    buf = token + "."
                else:
                    buf = candidate
            if buf:
                sentences.append(buf.strip())

            for s in sentences:
                if len(current) + len(s) + 1 <= max_chars:
                    current = (current + " " + s).strip()
                else:
                    flush_current()
                    current = s
    flush_current()
    return [b for b in blocks if b]


def synthesize_chunks_to_mp3(chunks: List[str], out_mp3: str, lang: str = "pt", slow: bool = False):
    """
    Converte blocos de texto em MP3 único:
    - Cada bloco vira um MP3 parcial via gTTS
    - Concatena tudo com pydub em um único arquivo final
    """
    if not chunks:
        raise ValueError("Documento vazio: não há texto para sintetizar.")

    with tempfile.TemporaryDirectory() as td:
        part_files = []
        for i, chunk in enumerate(chunks, 1):
            tts = gTTS(text=chunk, lang=lang, slow=slow)
            part_path = os.path.join(td, f"part_{i:03d}.mp3")
            tts.save(part_path)
            part_files.append(part_path)

        # Carrega e concatena
        merged = AudioSegment.empty()
        for pf in part_files:
            seg = AudioSegment.from_mp3(pf)
            merged += seg

        # Exporta MP3 final (bitrate padrão do pydub/ffmpeg; ajuste se quiser)
        merged.export(out_mp3, format="mp3")


def main():
    parser = argparse.ArgumentParser(description="Lê .docx e gera narração em MP3.")
    parser.add_argument("input_docx", help="Caminho para o arquivo .docx de entrada")
    parser.add_argument("output_mp3", help="Caminho para o arquivo .mp3 de saída")
    parser.add_argument("--lang", default="pt", help="Código de idioma (ex.: pt, pt-br, en, es)")
    parser.add_argument("--slow", default="false", choices=["true", "false"], help="Voz pausada")
    parser.add_argument("--max_chars", type=int, default=2000, help="Máximo de caracteres por bloco")
    args = parser.parse_args()

    if not os.path.exists(args.input_docx):
        raise FileNotFoundError(f"Arquivo não encontrado: {args.input_docx}")

    text = extract_text_from_docx(args.input_docx)
    chunks = sentencewise_chunks(text, max_chars=args.max_chars)
    if not chunks:
        raise ValueError("Não foi possível extrair texto utilizável do documento.")

    synthesize_chunks_to_mp3(
        chunks,
        args.output_mp3,
        lang=args.lang,
        slow=(args.slow.lower() == "true")
    )
    print(f"MP3 gerado com sucesso em: {args.output_mp3}")


if __name__ == "__main__":
    main()